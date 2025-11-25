"""Create a sample Word document for testing."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

doc = Document()

# Title
title = doc.add_heading("Sample Contract Agreement", 0)

# Section 1
doc.add_heading("1. Definitions", level=1)
doc.add_paragraph(
    "In this Agreement, the following terms shall have the meanings set out below:"
)

# Table with definitions
table = doc.add_table(rows=3, cols=2)
table.style = "Table Grid"

# Header row
table.rows[0].cells[0].text = "Term"
table.rows[0].cells[1].text = "Definition"

# Data rows
table.rows[1].cells[0].text = "Agreement"
table.rows[1].cells[1].text = "This contract and all schedules attached hereto."

table.rows[2].cells[0].text = "Services"
cell = table.rows[2].cells[1]
cell.text = "The professional services to be provided, including:"
# Add a paragraph within the cell
p = cell.add_paragraph("• Consulting services")
p = cell.add_paragraph("• Technical support")
p = cell.add_paragraph("• Training sessions")

# Section 2
doc.add_heading("2. Scope of Services", level=1)
doc.add_paragraph(
    "The Contractor agrees to provide the Services as described in Schedule A."
)

# Subsection 2a
doc.add_heading("2a. Primary Deliverables", level=2)
doc.add_paragraph("The following deliverables shall be provided:")

# Deliverables table
table2 = doc.add_table(rows=4, cols=3)
table2.style = "Table Grid"

table2.rows[0].cells[0].text = "Item"
table2.rows[0].cells[1].text = "Description"
table2.rows[0].cells[2].text = "Timeline"

table2.rows[1].cells[0].text = "Documentation"
table2.rows[1].cells[
    1
].text = "Complete technical documentation for all systems."
table2.rows[1].cells[2].text = "Week 2"

table2.rows[2].cells[0].text = "Training Materials"
table2.rows[2].cells[1].text = "User guides and training videos."
table2.rows[2].cells[2].text = "Week 4"

table2.rows[3].cells[0].text = "Final Report"
table2.rows[3].cells[
    1
].text = "Summary of all work completed, including recommendations for future improvements."
table2.rows[3].cells[2].text = "Week 6"

# Subsection 2b
doc.add_heading("2b. Secondary Deliverables", level=2)
doc.add_paragraph(
    "Additional deliverables may be requested as per Section 2a requirements."
)

# Section 3
doc.add_heading("3. Payment Terms", level=1)
para = doc.add_paragraph("Payment shall be made according to the schedule in ")
run = para.add_run("Schedule B")
run.bold = True
para.add_run(". See Section 2a for deliverable milestones.")

# Add some formatted text
doc.add_paragraph()
para = doc.add_paragraph()
para.add_run("Important: ").bold = True
para.add_run("All payments are due within ")
para.add_run("30 days").italic = True
para.add_run(" of invoice receipt.")

doc.save("sample_contract.docx")
print("Created sample_contract.docx")
