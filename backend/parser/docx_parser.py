"""
DOCX Parser - Extracts structured content from Word documents.

Handles:
- Paragraphs with text formatting (bold, italic)
- Numbered/bulleted lists with computed numbering
- Tables with rich cell content
- Headings
"""

import re
import uuid
from dataclasses import dataclass, field
from io import BytesIO
from typing import Optional

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


@dataclass
class TextRun:
    """A run of text with formatting."""

    text: str
    bold: bool = False
    italic: bool = False


@dataclass
class Paragraph:
    """A paragraph containing text runs."""

    runs: list[TextRun] = field(default_factory=list)
    style: Optional[str] = None
    numbering: Optional[str] = None  # Computed numbering like "2a"
    level: int = 0  # Heading level (0 = not a heading)


@dataclass
class TableCell:
    """A table cell containing block content."""

    content: list = field(
        default_factory=list
    )  # List of Paragraph or nested Table


@dataclass
class TableRow:
    """A table row containing cells."""

    cells: list[TableCell] = field(default_factory=list)


@dataclass
class Table:
    """A table with rows and cells."""

    id: str = field(default_factory=lambda: str(uuid.uuid4()))
    rows: list[TableRow] = field(default_factory=list)


@dataclass
class Section:
    """A document section with content and optional children."""

    id: str = field(default_factory=lambda: str(uuid.uuid4()))
    original_ref: Optional[str] = None
    level: int = 1
    title: str = ""
    content: list = field(default_factory=list)  # Paragraph, Table
    children: list["Section"] = field(default_factory=list)


class NumberingTracker:
    """Tracks list numbering state to compute actual numbers."""

    def __init__(self, document: Document):
        self.document = document
        self.counters: dict[str, list[int]] = (
            {}
        )  # numId -> [level0_count, level1_count, ...]
        self._numbering_formats = self._extract_numbering_formats()

    def _extract_numbering_formats(self) -> dict:
        """Extract numbering format definitions from document."""
        formats = {}
        try:
            numbering_part = self.document.part.numbering_part
        except (KeyError, NotImplementedError):
            # Document has no numbering definitions
            return formats
        if numbering_part is None:
            return formats

        # Parse the numbering definitions
        numbering_xml = numbering_part._element
        for abstract_num in numbering_xml.findall(qn("w:abstractNum")):
            abstract_id = abstract_num.get(qn("w:abstractNumId"))
            levels = {}
            for lvl in abstract_num.findall(qn("w:lvl")):
                ilvl = int(lvl.get(qn("w:ilvl")))
                num_fmt_elem = lvl.find(qn("w:numFmt"))
                lvl_text_elem = lvl.find(qn("w:lvlText"))

                num_fmt = (
                    num_fmt_elem.get(qn("w:val"))
                    if num_fmt_elem is not None
                    else "decimal"
                )
                lvl_text = (
                    lvl_text_elem.get(qn("w:val"))
                    if lvl_text_elem is not None
                    else "%1."
                )

                levels[ilvl] = {"format": num_fmt, "text": lvl_text}
            formats[abstract_id] = levels

        # Map numId to abstractNumId
        self._num_to_abstract = {}
        for num in numbering_xml.findall(qn("w:num")):
            num_id = num.get(qn("w:numId"))
            abstract_ref = num.find(qn("w:abstractNumId"))
            if abstract_ref is not None:
                self._num_to_abstract[num_id] = abstract_ref.get(qn("w:val"))

        return formats

    def get_number(self, num_id: str, ilvl: int) -> str:
        """Get the computed number for a list item."""
        if num_id not in self.counters:
            self.counters[num_id] = [0] * 10  # Support up to 10 levels

        # Increment current level, reset deeper levels
        self.counters[num_id][ilvl] += 1
        for i in range(ilvl + 1, 10):
            self.counters[num_id][i] = 0

        # Get format info
        abstract_id = self._num_to_abstract.get(num_id)
        if abstract_id and abstract_id in self._numbering_formats:
            level_info = self._numbering_formats[abstract_id].get(ilvl, {})
            num_fmt = level_info.get("format", "decimal")
            lvl_text = level_info.get("text", "%1.")
        else:
            num_fmt = "decimal"
            lvl_text = "%1."

        # Build the number string
        result = lvl_text
        for i in range(ilvl + 1):
            count = self.counters[num_id][i]
            formatted = self._format_number(
                count, num_fmt if i == ilvl else "decimal"
            )
            result = result.replace(f"%{i+1}", formatted)

        return result

    def _format_number(self, n: int, fmt: str) -> str:
        """Format a number according to the numbering format."""
        if fmt == "decimal":
            return str(n)
        elif fmt == "lowerLetter":
            return chr(ord("a") + n - 1) if 1 <= n <= 26 else str(n)
        elif fmt == "upperLetter":
            return chr(ord("A") + n - 1) if 1 <= n <= 26 else str(n)
        elif fmt == "lowerRoman":
            return self._to_roman(n).lower()
        elif fmt == "upperRoman":
            return self._to_roman(n)
        elif fmt == "bullet":
            return "•"
        else:
            return str(n)

    def _to_roman(self, n: int) -> str:
        """Convert integer to Roman numerals."""
        if n <= 0:
            return str(n)
        result = ""
        for value, numeral in [
            (1000, "M"),
            (900, "CM"),
            (500, "D"),
            (400, "CD"),
            (100, "C"),
            (90, "XC"),
            (50, "L"),
            (40, "XL"),
            (10, "X"),
            (9, "IX"),
            (5, "V"),
            (4, "IV"),
            (1, "I"),
        ]:
            while n >= value:
                result += numeral
                n -= value
        return result


def parse_paragraph(
    para, numbering_tracker: Optional[NumberingTracker] = None
) -> Paragraph:
    """Parse a python-docx paragraph into our intermediate format."""
    runs = []
    for run in para.runs:
        if run.text:
            runs.append(
                TextRun(
                    text=run.text,
                    bold=run.bold or False,
                    italic=run.italic or False,
                )
            )

    # Detect heading level
    level = 0
    style_name = para.style.name if para.style else ""
    if style_name.startswith("Heading"):
        try:
            level = int(style_name.replace("Heading ", ""))
        except ValueError:
            pass

    # Get numbering info
    numbering = None
    if numbering_tracker and para._element.pPr is not None:
        num_pr = para._element.pPr.find(qn("w:numPr"))
        if num_pr is not None:
            ilvl_elem = num_pr.find(qn("w:ilvl"))
            num_id_elem = num_pr.find(qn("w:numId"))
            if ilvl_elem is not None and num_id_elem is not None:
                ilvl = int(ilvl_elem.get(qn("w:val")))
                num_id = num_id_elem.get(qn("w:val"))
                if num_id != "0":  # numId 0 means no numbering
                    numbering = numbering_tracker.get_number(num_id, ilvl)

    return Paragraph(
        runs=runs, style=style_name, numbering=numbering, level=level
    )


def parse_table(
    table, numbering_tracker: Optional[NumberingTracker] = None
) -> Table:
    """Parse a python-docx table into our intermediate format."""
    parsed_table = Table()

    for row in table.rows:
        parsed_row = TableRow()
        for cell in row.cells:
            parsed_cell = TableCell()

            # Parse cell content - cells can contain paragraphs and nested tables
            for para in cell.paragraphs:
                parsed_para = parse_paragraph(para, numbering_tracker)
                if parsed_para.runs:  # Only add non-empty paragraphs
                    parsed_cell.content.append(parsed_para)

            # Check for nested tables
            for nested_table in cell.tables:
                parsed_cell.content.append(
                    parse_table(nested_table, numbering_tracker)
                )

            # Ensure cell has at least one paragraph (empty cell)
            if not parsed_cell.content:
                parsed_cell.content.append(Paragraph(runs=[TextRun(text="")]))

            parsed_row.cells.append(parsed_cell)

        parsed_table.rows.append(parsed_row)

    return parsed_table


def parse_docx(file_content: bytes) -> list:
    """
    Parse a DOCX file and return a list of document elements.

    Args:
        file_content: Raw bytes of the DOCX file

    Returns:
        List of Paragraph, Table, and Section objects
    """
    doc = Document(BytesIO(file_content))
    numbering_tracker = NumberingTracker(doc)

    elements = []

    # Build lookup maps for paragraphs and tables by their XML element
    para_map = {p._element: p for p in doc.paragraphs}
    table_map = {t._tbl: t for t in doc.tables}

    # Iterate through document body in order
    for element in doc.element.body:
        if element.tag == qn("w:p"):
            # It's a paragraph
            if element in para_map:
                p = para_map[element]
                parsed = parse_paragraph(p, numbering_tracker)
                if parsed.runs or parsed.numbering:
                    elements.append(parsed)
        elif element.tag == qn("w:tbl"):
            # It's a table
            if element in table_map:
                t = table_map[element]
                elements.append(parse_table(t, numbering_tracker))

    return elements


def elements_to_dict(elements: list) -> list[dict]:
    """Convert parsed elements to JSON-serializable dictionaries."""
    result = []

    for elem in elements:
        if isinstance(elem, Paragraph):
            result.append(
                {
                    "type": "paragraph",
                    "runs": [
                        {"text": r.text, "bold": r.bold, "italic": r.italic}
                        for r in elem.runs
                    ],
                    "style": elem.style,
                    "numbering": elem.numbering,
                    "level": elem.level,
                }
            )
        elif isinstance(elem, Table):
            result.append(
                {
                    "type": "table",
                    "id": elem.id,
                    "rows": [
                        {
                            "cells": [
                                {"content": elements_to_dict(cell.content)}
                                for cell in row.cells
                            ]
                        }
                        for row in elem.rows
                    ],
                }
            )
        elif isinstance(elem, Section):
            result.append(
                {
                    "type": "section",
                    "id": elem.id,
                    "originalRef": elem.original_ref,
                    "level": elem.level,
                    "title": elem.title,
                    "content": elements_to_dict(elem.content),
                    "children": elements_to_dict(elem.children),
                }
            )

    return result
