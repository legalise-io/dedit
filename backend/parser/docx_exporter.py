"""
DOCX Exporter - Converts TipTap JSON back to Word document format.

This module handles the reverse conversion from TipTap editor JSON
back to a Word document (.docx) using python-docx.
"""

from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def create_docx_from_tiptap(tiptap_json: dict) -> BytesIO:
    """
    Convert TipTap JSON document to a Word document.

    Args:
        tiptap_json: TipTap document JSON with structure:
            {
                "type": "doc",
                "content": [...]
            }

    Returns:
        BytesIO buffer containing the .docx file
    """
    doc = Document()

    content = tiptap_json.get("content", [])
    for node in content:
        process_node(doc, node)

    # Save to BytesIO buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def process_node(doc: Document, node: dict, table_cell=None) -> None:
    """
    Process a TipTap node and add it to the document.

    Args:
        doc: The python-docx Document object
        node: A TipTap node dictionary
        table_cell: Optional table cell to add content to (for nested content)
    """
    node_type = node.get("type")

    if node_type == "paragraph":
        process_paragraph(doc, node, table_cell)
    elif node_type == "heading":
        process_heading(doc, node, table_cell)
    elif node_type == "table":
        process_table(doc, node)
    elif node_type == "section":
        process_section(doc, node)


def process_paragraph(doc: Document, node: dict, table_cell=None) -> None:
    """
    Process a paragraph node.

    Args:
        doc: The python-docx Document object
        node: The paragraph node
        table_cell: Optional table cell to add paragraph to
    """
    if table_cell is not None:
        para = (
            table_cell.paragraphs[0]
            if table_cell.paragraphs
            else table_cell.add_paragraph()
        )
        # Clear default paragraph if it exists and is empty
        if para.text == "" and len(table_cell.paragraphs) == 1:
            pass  # Use existing empty paragraph
        else:
            para = table_cell.add_paragraph()
    else:
        para = doc.add_paragraph()

    content = node.get("content", [])
    for text_node in content:
        if text_node.get("type") == "text":
            add_text_run(para, text_node)


def process_heading(doc: Document, node: dict, table_cell=None) -> None:
    """
    Process a heading node.

    Args:
        doc: The python-docx Document object
        node: The heading node
        table_cell: Optional table cell (headings in cells become bold paragraphs)
    """
    level = node.get("attrs", {}).get("level", 1)
    content = node.get("content", [])

    if table_cell is not None:
        # In a table cell, just make it a bold paragraph
        para = table_cell.add_paragraph()
        for text_node in content:
            if text_node.get("type") == "text":
                run = para.add_run(text_node.get("text", ""))
                run.bold = True
                apply_marks(run, text_node.get("marks", []))
    else:
        # Use Word's heading styles
        heading_style = f"Heading {min(level, 9)}"
        para = doc.add_heading(level=level)

        for text_node in content:
            if text_node.get("type") == "text":
                add_text_run(para, text_node)


def process_table(doc: Document, node: dict) -> None:
    """
    Process a table node.

    Args:
        doc: The python-docx Document object
        node: The table node
    """
    rows_data = node.get("content", [])
    if not rows_data:
        return

    # Determine table dimensions
    num_rows = len(rows_data)
    num_cols = (
        max(len(row.get("content", [])) for row in rows_data)
        if rows_data
        else 0
    )

    if num_rows == 0 or num_cols == 0:
        return

    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = "Table Grid"

    for row_idx, row_node in enumerate(rows_data):
        cells = row_node.get("content", [])
        for col_idx, cell_node in enumerate(cells):
            if col_idx < num_cols:
                cell = table.rows[row_idx].cells[col_idx]
                # Clear default paragraph
                if cell.paragraphs:
                    cell.paragraphs[0].clear()

                # Process cell content
                cell_content = cell_node.get("content", [])
                for i, content_node in enumerate(cell_content):
                    if i == 0 and cell.paragraphs:
                        # Use existing first paragraph
                        if content_node.get("type") == "paragraph":
                            for text_node in content_node.get("content", []):
                                if text_node.get("type") == "text":
                                    add_text_run(cell.paragraphs[0], text_node)
                        elif content_node.get("type") == "heading":
                            for text_node in content_node.get("content", []):
                                if text_node.get("type") == "text":
                                    run = cell.paragraphs[0].add_run(
                                        text_node.get("text", "")
                                    )
                                    run.bold = True
                                    apply_marks(run, text_node.get("marks", []))
                    else:
                        process_node(doc, content_node, table_cell=cell)


def process_section(doc: Document, node: dict) -> None:
    """
    Process a section node (just process its content).

    Args:
        doc: The python-docx Document object
        node: The section node
    """
    content = node.get("content", [])
    for child_node in content:
        process_node(doc, child_node)


def add_text_run(para, text_node: dict) -> None:
    """
    Add a text run to a paragraph with formatting.

    Args:
        para: The python-docx Paragraph object
        text_node: The TipTap text node
    """
    text = text_node.get("text", "")
    if not text:
        return

    run = para.add_run(text)
    marks = text_node.get("marks", [])
    apply_marks(run, marks)


def apply_marks(run, marks: list) -> None:
    """
    Apply TipTap marks (formatting) to a Word run.

    Args:
        run: The python-docx Run object
        marks: List of TipTap mark dictionaries
    """
    for mark in marks:
        mark_type = mark.get("type")

        if mark_type == "bold":
            run.bold = True
        elif mark_type == "italic":
            run.italic = True
        # insertion and deletion marks are ignored - the text is already
        # in its final state after accept/reject operations
        # comment marks are also ignored in export
